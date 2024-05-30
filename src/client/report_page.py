import flet
from src.database.models import Report
from docx import Document


class ReportPage(flet.SafeArea):

    def get_reports(self) -> list[flet.Container]:
        def on_save_click(event: flet.ControlEvent):
            def on_file_pick(e: flet.FilePickerResultEvent):
                report_id = event.control.data
                report = Report.get(Report.id == report_id)
                self.generate_word_document(report, e.path)

            file_picker = flet.FilePicker(
                on_result=on_file_pick
            )

            self.page.overlay.append(
                file_picker
            ),

            self.page.update()
            self.page.overlay[-1].save_file(
                file_name=Report.get(Report.id==event.control.data).from_user.fullname,
                file_type='docx'
            )

        return [
            flet.Container(
                content=flet.Row(
                    alignment=flet.MainAxisAlignment.SPACE_BETWEEN,
                    controls=[
                        flet.Text(report.title + ' от ' + str(report.created_date)),
                        flet.TextButton(
                            text='Сохранить',
                            icon=flet.icons.SAVE,
                            on_click=on_save_click,
                            data=report.id
                        )
                    ]
                )
            )
            for report in Report.select()
        ]

    def generate_word_document(self, report, path):
        doc = Document()
        doc.add_heading(f'{report.title} от {report.created_date}', 0)
        doc.add_heading(f'Автор: {report.from_user.fullname}', level=1)
        doc.add_paragraph(report.text)

        doc.save(path)

    def on_update_click(self, event: flet.ControlEvent):
        self.content.controls[1].controls = self.get_reports()
        self.update()

    def build(self):
        self.expand = True
        self.content = flet.Column(
            horizontal_alignment=flet.CrossAxisAlignment.CENTER,
            scroll=flet.ScrollMode.AUTO,
            controls=[
                flet.TextButton(
                    icon=flet.icons.AUTORENEW,
                    text='Обновить',
                    on_click=self.on_update_click
                ),

                flet.Column(
                    controls=self.get_reports()
                )
            ]
        )
